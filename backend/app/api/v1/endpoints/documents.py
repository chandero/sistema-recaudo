"""
Endpoints para generación de documentos y gestión de correspondencia
"""
import traceback
import os
import shutil
import json
import tempfile
import glob
import re
import zipfile
import subprocess
from copy import deepcopy
from datetime import datetime
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, Body
from fastapi.responses import FileResponse
from sqlmodel import Session, select
from pydantic import BaseModel
from num2words import num2words
from docx import Document as DocxDocument

from app.core.database import get_session
from app.core.dependencies import get_current_active_user
from app.models.user import User
from app.models.document import DocumentTemplate, GeneratedDocument
from app.schemas.document import (
    DocumentTemplateCreate, 
    DocumentTemplateResponse, 
    DocumentGenerationRequest,
    BatchGenerationRequest,
    GeneratedDocumentResponse
)
from app.models.process import CobroProcess
from app.models.client import Client
from app.models.obligation import Obligation
from app.services.document_generation_service import DocumentGenerationService

router = APIRouter()

UPLOAD_DIR = "app/uploads/templates"
OUTPUT_DIR = "app/uploads/generated"

# Asegurar directorios
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)


class CorrespondenceBatchRequest(BaseModel):
    obligation_ids: List[int]
    template_id: Optional[int] = None


class CorrespondencePrintRangeRequest(BaseModel):
    resolution_from: int
    resolution_to: int
    template_id: Optional[int] = None
    output_format: str = "docx"


def _safe_filename(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", str(value or "sin_dato")).strip("_")


def _amount_in_words(value: float) -> str:
    """Representación del importe en letras para el acto administrativo."""
    integer = int(round(value or 0))
    return f"{num2words(integer, lang='es').upper()} PESOS M/CTE"


def _correspondence_context(obligation: Obligation, client: Client) -> dict:
    resolution_date = obligation.resolution_date.strftime("%d/%m/%Y")
    cutoff_date = obligation.due_date.strftime("%d/%m/%Y") if obligation.due_date else resolution_date
    return {
        "numero_resolucion": obligation.resolution_number,
        "anho_resolucion": obligation.resolution_year or obligation.resolution_date.year,
        "fecha_resolucion": resolution_date,
        "nombre_usuario": client.name,
        "documento_usuario": client.identification,
        "direccion_usuario": client.address or "",
        "cuenta_usuario": obligation.description or str(obligation.id),
        "deuda_numero": f"{obligation.amount:,.2f}",
        "deuda_letras": _amount_in_words(obligation.amount),
        "fecha_corte": cutoff_date,
    }


@router.post("/correspondence/print-range")
def generate_correspondence_print_range(
    request: CorrespondencePrintRangeRequest,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_active_user),
):
    """Genera un único DOCX imprimible para un rango de resoluciones."""
    if request.resolution_from > request.resolution_to:
        raise HTTPException(status_code=400, detail="El número inicial no puede superar el número final")
    if request.output_format not in {"docx", "pdf"}:
        raise HTTPException(status_code=400, detail="El formato debe ser docx o pdf")

    template_statement = select(DocumentTemplate).where(
        DocumentTemplate.tenant_id == current_user.tenant_id,
        DocumentTemplate.is_active == True,
    )
    template_statement = (
        template_statement.where(DocumentTemplate.id == request.template_id)
        if request.template_id
        else template_statement.where(DocumentTemplate.code == "GD-F.03")
    )
    template = session.exec(template_statement).first()
    if not template:
        raise HTTPException(status_code=404, detail="No existe una plantilla activa GD-F.03")

    tenant_rows = session.exec(
        select(Obligation, Client)
        .join(Client, Obligation.client_id == Client.id)
        .where(Client.tenant_id == current_user.tenant_id)
    ).all()
    rows = [
        (obligation, client)
        for obligation, client in tenant_rows
        if obligation.resolution_number
        and obligation.resolution_number.isdigit()
        and request.resolution_from <= int(obligation.resolution_number) <= request.resolution_to
    ]
    rows.sort(key=lambda row: int(row[0].resolution_number))
    if not rows:
        raise HTTPException(status_code=404, detail="No hay correspondencia en el rango indicado")
    if any(not obligation.resolution_date for obligation, _ in rows):
        raise HTTPException(status_code=400, detail="El rango contiene obligaciones sin fecha de resolución")

    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    batch_dir = os.path.join(OUTPUT_DIR, f"impresion_{current_user.tenant_id}_{stamp}")
    os.makedirs(batch_dir, exist_ok=True)
    template_path = os.path.join(batch_dir, "template.docx")
    with open(template_path, "wb") as template_file:
        template_file.write(template.content)

    rendered_paths = []
    for obligation, client in rows:
        output_path = os.path.join(batch_dir, f"resolucion_{obligation.resolution_number}.docx")
        DocumentGenerationService.render_template(
            template_path,
            _correspondence_context(obligation, client),
            output_path,
        )
        rendered_paths.append(output_path)

    master = DocxDocument(rendered_paths[0])
    for path in rendered_paths[1:]:
        master.add_page_break()
        source = DocxDocument(path)
        for element in source.element.body:
            if not element.tag.endswith("}sectPr"):
                master.element.body.append(deepcopy(element))
    output_path = os.path.join(
        OUTPUT_DIR,
        f"impresion_correspondencia_{request.resolution_from}_{request.resolution_to}_{stamp}.docx",
    )
    master.save(output_path)
    if request.output_format == "pdf":
        profile_dir = os.path.join(batch_dir, "libreoffice-profile")
        os.makedirs(profile_dir, exist_ok=True)
        conversion = subprocess.run(
            [
                "libreoffice",
                "--headless",
                f"-env:UserInstallation=file://{os.path.abspath(profile_dir)}",
                "--convert-to",
                "pdf",
                "--outdir",
                OUTPUT_DIR,
                output_path,
            ],
            capture_output=True,
            text=True,
            # Un rango completo puede contener miles de páginas. LibreOffice
            # necesita más de tres minutos para convertir lotes grandes.
            timeout=1800,
        )
        pdf_path = os.path.splitext(output_path)[0] + ".pdf"
        if conversion.returncode != 0 or not os.path.exists(pdf_path):
            raise HTTPException(
                status_code=500,
                detail=f"No se pudo convertir el documento a PDF: {conversion.stderr.strip()}",
            )
        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename=os.path.basename(pdf_path),
            headers={"X-Document-Count": str(len(rows))},
        )
    return FileResponse(
        output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=os.path.basename(output_path),
        headers={"X-Document-Count": str(len(rows))},
    )


@router.post("/correspondence/batch")
def generate_certified_correspondence_batch(
    request: CorrespondenceBatchRequest,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_active_user),
):
    """Genera un ZIP con un DOCX individual por obligación resuelta."""
    if not request.obligation_ids:
        raise HTTPException(status_code=400, detail="Seleccione al menos una obligación")

    template_statement = select(DocumentTemplate).where(
        DocumentTemplate.tenant_id == current_user.tenant_id,
        DocumentTemplate.is_active == True,
    )
    if request.template_id:
        template_statement = template_statement.where(DocumentTemplate.id == request.template_id)
    else:
        template_statement = template_statement.where(DocumentTemplate.code == "GD-F.03")
    template = session.exec(template_statement).first()
    if not template:
        raise HTTPException(status_code=404, detail="No existe una plantilla activa GD-F.03")

    rows = session.exec(
        select(Obligation, Client)
        .join(Client, Obligation.client_id == Client.id)
        .where(
            Obligation.id.in_(request.obligation_ids),
            Client.tenant_id == current_user.tenant_id,
        )
        .order_by(Obligation.resolution_number)
    ).all()
    if len(rows) != len(set(request.obligation_ids)):
        raise HTTPException(status_code=400, detail="Algunas obligaciones no existen o no pertenecen al tenant")

    missing = [obligation.id for obligation, _ in rows if not obligation.resolution_number or not obligation.resolution_date]
    if missing:
        raise HTTPException(status_code=400, detail=f"Hay {len(missing)} obligaciones sin resolución o fecha")

    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    batch_dir = os.path.join(OUTPUT_DIR, f"correspondencia_{current_user.tenant_id}_{stamp}")
    os.makedirs(batch_dir, exist_ok=True)
    template_path = os.path.join(batch_dir, "template.docx")
    with open(template_path, "wb") as template_file:
        template_file.write(template.content)

    zip_path = os.path.join(OUTPUT_DIR, f"correspondencia_certificada_{stamp}.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for obligation, client in rows:
            context = _correspondence_context(obligation, client)
            filename = (
                f"resolucion_{_safe_filename(obligation.resolution_number)}_"
                f"{_safe_filename(client.identification)}.docx"
            )
            output_path = os.path.join(batch_dir, filename)
            DocumentGenerationService.render_template(template_path, context, output_path)
            zip_file.write(output_path, filename)

    return FileResponse(
        zip_path,
        media_type="application/zip",
        filename=os.path.basename(zip_path),
        headers={"X-Document-Count": str(len(rows))},
    )


@router.get("/templates", response_model=List[DocumentTemplateResponse])
def get_templates(
    skip: int = 0,
    limit: int = 100,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_active_user)
):
    """Obtener plantillas del tenant"""
    try:
        statement = select(DocumentTemplate).where(
            DocumentTemplate.tenant_id == current_user.tenant_id,
            DocumentTemplate.is_active == True
        ).offset(skip).limit(limit)
        templates = session.exec(statement).all()
        return templates
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo plantillas: {str(e)}")


@router.post("/templates", response_model=DocumentTemplateResponse)
def create_template(
    name: str = Form(...),
    code: str = Form(...),
    description: Optional[str] = Form(None),
    variables_schema: str = Form("{}"),
    file: UploadFile = File(...),
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_active_user)
):
    """Subir nueva plantilla documental"""
    try:
        # Validar unicidad de código
        statement = select(DocumentTemplate).where(
            DocumentTemplate.code == code,
            DocumentTemplate.tenant_id == current_user.tenant_id
        )
        existing = session.exec(statement).first()
        if existing:
            raise HTTPException(status_code=400, detail="El código de plantilla ya existe")
        
        # Validar tipo de archivo
        allowed_extensions = [".docx", ".dotx"]
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in allowed_extensions:
            raise HTTPException(status_code=400, detail=f"Tipo de archivo no permitido. Solo se permiten: {', '.join(allowed_extensions)}")
        
        # Guardar archivo
        file_path = os.path.join(UPLOAD_DIR, f"{current_user.tenant_id}_{file.filename}")
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Leer contenido del archivo
        with open(file_path, "rb") as f:
            file_content = f.read()
        
        # Parse variables_schema if it's a string
        if isinstance(variables_schema, str):
            try:
                variables_schema = json.loads(variables_schema)
            except (json.JSONDecodeError, TypeError):
                variables_schema = {}
        
        # Crear registro
        template = DocumentTemplate(
            name=name,
            code=code,
            description=description,
            content=file_content,
            file_path=file_path,  # Añadir esta propiedad al modelo si es necesario
            variables_schema=variables_schema,
            tenant_id=current_user.tenant_id,
            version=1
        )
        session.add(template)
        session.commit()
        session.refresh(template)
        return template
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creando plantilla: {str(e)}")


@router.get("/generated", response_model=List[GeneratedDocumentResponse])
def get_generated_documents(
    skip: int = 0,
    limit: int = 100,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_active_user)
):
    """Obtener documentos generados del tenant"""
    try:
        statement = select(GeneratedDocument).where(
            GeneratedDocument.tenant_id == current_user.tenant_id
        ).offset(skip).limit(limit)
        documents = session.exec(statement).all()
        return documents
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo documentos generados: {str(e)}")


@router.post("/generate")
def generate_document(
    request: DocumentGenerationRequest,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_active_user)
):
    """Generar un documento individual"""
    try:
        statement = select(DocumentTemplate).where(
            DocumentTemplate.id == request.template_id,
            DocumentTemplate.tenant_id == current_user.tenant_id
        )
        template = session.exec(statement).first()
        if not template:
            raise HTTPException(status_code=404, detail="Plantilla no encontrada")
        
        output_filename = f"doc_{request.process_id}_{template.code}.pdf"
        output_path = os.path.join(OUTPUT_DIR, output_filename)
        
        # Crear archivo temporal DOCX
        temp_docx = output_path.replace(".pdf", ".docx")
        
        # Generar documento
        DocumentGenerationService.render_template(
            template.file_path, 
            request.variables, 
            temp_docx
        )
        
        # Convertir a PDF
        DocumentGenerationService.convert_to_pdf(
            temp_docx,
            output_path
        )
        
        # Registrar en BD
        doc_record = GeneratedDocument(
            process_id=request.process_id,
            template_id=template.id,
            file_path=output_path,
            filename=output_filename,
            document_type="PDF",
            tenant_id=current_user.tenant_id,
            created_by=current_user.id
        )
        session.add(doc_record)
        session.commit()
        
        # Limpieza
        if os.path.exists(temp_docx):
            os.remove(temp_docx)
        
        return {"message": "Documento generado", "path": output_path}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando documento: {str(e)}")


@router.post("/templates/{template_id}/generate")
def generate_document_from_template(
    template_id: int,
    variables: dict = Body(...),
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_active_user)
):
    """Generar un documento directamente desde una plantilla con variables"""
    from app.core.logging.logger import get_logger
    logger = get_logger(__name__)
    
    try:
        logger.info(
            module="document_generation",
            action="generate_template_document_start",
            message=f"Iniciando generación de documento desde plantilla ID: {template_id}, Usuario: {current_user.id}, Tenant: {current_user.tenant_id}",
            data={"template_id": template_id, "user_id": current_user.id, "tenant_id": current_user.tenant_id}
        )
        
        statement = select(DocumentTemplate).where(
            DocumentTemplate.id == template_id,
            DocumentTemplate.tenant_id == current_user.tenant_id
        )
        template = session.exec(statement).first()
        
        if not template:
            logger.warning(
                module="document_generation", 
                action="template_not_found",
                message=f"Plantilla no encontrada - ID: {template_id}, Tenant: {current_user.tenant_id}",
                data={"template_id": template_id, "tenant_id": current_user.tenant_id}
            )
            raise HTTPException(status_code=404, detail="Plantilla no encontrada")
        
        # El archivo se guarda con el formato: {tenant_id}_{nombre_original_del_archivo}
        # Necesitamos encontrar el archivo correcto basado en el código de la plantilla
        # Busquemos archivos que comiencen con el prefijo del tenant y contengan el código
        pattern = os.path.join(UPLOAD_DIR, f"{current_user.tenant_id}_*{template.code}*")
        matching_files = glob.glob(pattern)
        
        file_path = None
        if matching_files:
            # Tomar el primer archivo que coincida
            file_path = matching_files[0]
        else:
            # Si no se encuentra con el patrón anterior, probar con el nombre exacto del código
            file_extension = ".docx" if template.template_type == "DOCX" else ".dotx"
            file_path = os.path.join(UPLOAD_DIR, f"{current_user.tenant_id}_{template.code}{file_extension}")
        
        logger.info(
            module="document_generation",
            action="template_found",
            message=f"Plantilla encontrada: {template.name}, Ruta: {file_path}",
            data={"template_id": template.id, "template_name": template.name, "file_path": file_path}
        )
        
        # Verificar que el archivo de plantilla exista
        if not os.path.exists(file_path):
            logger.error(
                module="document_generation",
                action="template_file_not_found",
                message=f"Archivo de plantilla no encontrado en: {file_path}",
                data={"template_id": template.id, "file_path": file_path}
            )
            # Intentar encontrar el archivo con otros posibles nombres
            # Buscar archivos que comiencen con el tenant_id y contengan el código
            possible_patterns = [
                os.path.join(UPLOAD_DIR, f"{current_user.tenant_id}_{template.code}.*"),
                os.path.join(UPLOAD_DIR, f"{current_user.tenant_id}_{template.code}*"),
                os.path.join(UPLOAD_DIR, f"*{template.code}*")
            ]
            
            found_file = None
            for pattern in possible_patterns:
                matching_files = glob.glob(pattern)
                if matching_files:
                    found_file = matching_files[0]
                    break
            
            if found_file:
                file_path = found_file
                logger.info(
                    module="document_generation",
                    action="template_file_found_alternative",
                    message=f"Archivo encontrado con nombre alternativo: {found_file}",
                    data={"template_id": template.id, "file_path": found_file}
                )
            else:
                raise HTTPException(status_code=500, detail=f"Archivo de plantilla no encontrado en: {file_path}")
        
        logger.info(
            module="document_generation",
            action="template_file_exists",
            message=f"Archivo de plantilla existe y es accesible: {file_path}",
            data={"template_id": template.id, "file_path": file_path, "size": os.path.getsize(file_path)}
        )
        
        # Verificar que el archivo sea accesible
        if not os.access(file_path, os.R_OK):
            logger.error(
                module="document_generation",
                action="template_file_not_accessible",
                message=f"No se puede acceder al archivo de plantilla: {file_path}",
                data={"template_id": template.id, "file_path": file_path}
            )
            raise HTTPException(status_code=500, detail=f"No se puede acceder al archivo de plantilla: {file_path}")
        
        # Crear archivo temporal DOCX para vista previa
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_path = temp_file.name
        
        logger.info(
            module="document_generation",
            action="temp_file_created",
            message=f"Archivo temporal creado: {temp_path}",
            data={"temp_path": temp_path}
        )
        
        try:
            logger.info(
                module="document_generation",
                action="template_render_start",
                message=f"Iniciando renderizado de plantilla: {file_path}",
                data={"template_id": template.id, "template_path": file_path, "variables_count": len(variables) if isinstance(variables, dict) else 0}
            )
            
            # Generar documento
            DocumentGenerationService.render_template(
                file_path, 
                variables, 
                temp_path
            )
            
            logger.info(
                module="document_generation",
                action="template_render_success",
                message=f"Plantilla renderizada exitosamente en: {temp_path}",
                data={"template_id": template.id, "output_path": temp_path}
            )
        except Exception as render_error:
            logger.error(
                module="document_generation",
                action="template_render_failed",
                message=f"Error al renderizar la plantilla: {str(render_error)}",
                data={"template_id": template.id, "template_path": file_path, "error": str(render_error)}
            )
            import traceback
            logger.error(
                module="document_generation",
                action="template_render_traceback",
                message=f"Detalle del error: {traceback.format_exc()}",
                data={"template_id": template.id, "error_details": traceback.format_exc()}
            )
            # Eliminar el archivo temporal si ocurrió un error en el renderizado
            if os.path.exists(temp_path):
                os.unlink(temp_path)
            raise HTTPException(status_code=500, detail=f"Error al renderizar la plantilla: {str(render_error)}")
        
        # Leer el contenido del archivo generado
        try:
            logger.info(
                module="document_generation",
                action="reading_generated_file",
                message=f"Leyendo archivo generado: {temp_path}",
                data={"output_path": temp_path}
            )
            with open(temp_path, "rb") as f:
                content = f.read()
            logger.info(
                module="document_generation",
                action="file_read_success",
                message=f"Archivo leído exitosamente, tamaño: {len(content)} bytes",
                data={"output_path": temp_path, "size": len(content)}
            )
        except Exception as read_error:
            logger.error(
                module="document_generation",
                action="file_read_failed",
                message=f"Error al leer el archivo generado: {str(read_error)}",
                data={"output_path": temp_path, "error": str(read_error)}
            )
            import traceback
            logger.error(
                module="document_generation",
                action="file_read_traceback",
                message=f"Detalle del error: {traceback.format_exc()}",
                data={"output_path": temp_path, "error_details": traceback.format_exc()}
            )
            if os.path.exists(temp_path):
                os.unlink(temp_path)
            raise HTTPException(status_code=500, detail=f"Error al leer el archivo generado: {str(read_error)}")
        
        # Eliminar el archivo temporal
        try:
            os.unlink(temp_path)
            logger.info(
                module="document_generation",
                action="temp_file_deleted",
                message=f"Archivo temporal eliminado: {temp_path}",
                data={"temp_path": temp_path}
            )
        except Exception as cleanup_error:
            logger.warning(
                module="document_generation",
                action="temp_file_delete_failed",
                message=f"Advertencia: Error al eliminar archivo temporal: {str(cleanup_error)}",
                data={"temp_path": temp_path, "error": str(cleanup_error)}
            )
        
        # Devolver el contenido del archivo generado como respuesta binaria
        from fastapi.responses import Response
        
        response = Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={template.code}_preview.docx"
            }
        )
        
        logger.info(
            module="document_generation",
            action="document_generation_success",
            message=f"Respuesta generada exitosamente para la plantilla: {template.name}",
            data={"template_id": template.id, "template_name": template.name}
        )
        return response
    except HTTPException:
        logger.error(
            module="document_generation",
            action="http_exception_occurred",
            message="HTTPException en generación de documento",
            data={"error_type": "HTTPException"}
        )
        import traceback
        logger.error(
            module="document_generation",
            action="http_exception_traceback",
            message=f"Detalle del error: {traceback.format_exc()}",
            data={"error_details": traceback.format_exc()}
        )
        raise
    except Exception as e:
        logger.error(
            module="document_generation",
            action="general_exception_occurred",
            message=f"Error general en generación de documento: {str(e)}",
            data={"error": str(e), "error_type": type(e).__name__}
        )
        import traceback
        logger.error(
            module="document_generation",
            action="general_exception_traceback",
            message=f"Detalle del error: {traceback.format_exc()}",
            data={"error_details": traceback.format_exc()}
        )
        raise HTTPException(status_code=500, detail=f"Error generando documento desde plantilla: {str(e)}")


@router.post("/generate/batch")
def generate_batch_documents(
    request: BatchGenerationRequest,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_active_user)
):
    """Generar lote masivo de documentos"""
    try:
        statement = select(DocumentTemplate).where(
            DocumentTemplate.id == request.template_id,
            DocumentTemplate.tenant_id == current_user.tenant_id
        )
        template = session.exec(statement).first()
        if not template:
            raise HTTPException(status_code=404, detail="Plantilla no encontrada")
        
        # Obtener datos de los procesos
        statement = select(CobroProcess).where(
            CobroProcess.id.in_(request.process_ids),
            CobroProcess.tenant_id == current_user.tenant_id
        )
        processes = session.exec(statement).all()
        
        if len(processes) != len(request.process_ids):
            raise HTTPException(status_code=400, detail="Algunos procesos no existen o no pertenecen al tenant")
        
        # Preparar datos para generación
        proc_data = []
        for p in processes:
            # Preparar datos del proceso para la plantilla
            process_data = {
                'id': p.id,
                'client_name': p.client.name if p.client else 'Cliente desconocido',
                'client_email': p.client.email if p.client and hasattr(p.client, 'email') else '',
                'client_phone': p.client.phone if p.client and hasattr(p.client, 'phone') else '',
                'amount': p.total_amount,
                'status': p.status.value if hasattr(p.status, 'value') else str(p.status),
                'created_at': p.created_at.isoformat() if p.created_at else '',
                'due_date': p.due_date.isoformat() if p.due_date else ''
            }
            proc_data.append(process_data)
        
        # Generar documentos
        output_dir = OUTPUT_DIR
        output_zip = DocumentGenerationService.generate_batch(
            [{'path': template.file_path}],
            proc_data,
            output_dir
        )
        
        return {"message": f"Lote de {len(processes)} documentos generado", "path": output_zip}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando lote de documentos: {str(e)}")
