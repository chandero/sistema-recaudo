"""
Servicio de generación de documentos para el sistema de recaudo.
Este servicio se encarga de la generación de documentos a partir de plantillas.
"""
import os
import re
import json
from pathlib import Path
from typing import Dict, List, Any
from docxtpl import DocxTemplate
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
import pandas as pd
from io import BytesIO
import zipfile
import tempfile


class DocumentGenerationService:
    """
    Servicio para la generación de documentos a partir de plantillas.
    Proporciona métodos para renderizar plantillas DOCX con datos,
    convertir a PDF y generar lotes de documentos.
    """

    @staticmethod
    def render_template(template_path: str, data: Dict[str, Any], output_path: str):
        """
        Renderiza una plantilla DOCX con los datos proporcionados.
        
        Args:
            template_path: Ruta a la plantilla DOCX
            data: Datos para rellenar la plantilla
            output_path: Ruta donde guardar el documento generado
        """
        try:
            # Cargar la plantilla
            doc = DocxTemplate(template_path)
            
            # Renderizar la plantilla con los datos
            doc.render(data)

            # La plantilla de resoluciones incluye un párrafo vacío justo
            # antes de "Proyecto:". LibreOffice le asigna altura de línea y
            # desplaza la última firma a una tercera página. Se elimina solo
            # cuando está realmente vacío y no contiene imágenes.
            paragraphs = doc.docx.paragraphs
            for index, paragraph in enumerate(paragraphs):
                if paragraph.text.strip().casefold() != "proyecto:":
                    continue
                previous = paragraph._p.getprevious()
                if (
                    previous is not None
                    and not "".join(previous.itertext()).strip()
                    and not previous.findall(f".//{qn('w:drawing')}")
                ):
                    previous.getparent().remove(previous)
                # El párrafo anterior ancla la imagen de la firma. No contiene
                # texto, pero LibreOffice en producción le reserva una línea
                # completa adicional. La imagen es flotante, por lo que esa
                # altura puede reducirse sin cambiar su tamaño ni posición.
                anchor_element = paragraph._p.getprevious()
                if (
                    anchor_element is not None
                    and anchor_element.findall(f".//{qn('w:drawing')}")
                ):
                    anchor_paragraph = next(
                        (
                            candidate
                            for candidate in paragraphs
                            if candidate._p is anchor_element
                        ),
                        None,
                    )
                    if anchor_paragraph is not None:
                        anchor_paragraph.paragraph_format.space_before = Pt(0)
                        anchor_paragraph.paragraph_format.space_after = Pt(0)
                        anchor_paragraph.paragraph_format.line_spacing = Pt(1)
                # Compactar exclusivamente las tres líneas finales de firma.
                # Los datos variables extensos pueden consumir unos puntos más
                # en la segunda página y desplazar la última línea a una tercera.
                for signature_paragraph in paragraphs[index:index + 3]:
                    signature_paragraph.paragraph_format.space_before = Pt(0)
                    signature_paragraph.paragraph_format.space_after = Pt(0)
                    signature_paragraph.paragraph_format.line_spacing = 1.0
                break
            
            # Guardar el documento generado
            doc.save(output_path)
        except Exception as e:
            raise Exception(f"Error al renderizar la plantilla: {str(e)}")

    @staticmethod
    def _restart_section_page_numbering(section_properties):
        """Hace que la numeración de una sección comience nuevamente en 1."""
        page_number_type = section_properties.find(qn("w:pgNumType"))
        if page_number_type is None:
            page_number_type = OxmlElement("w:pgNumType")
            section_properties.append(page_number_type)
        page_number_type.set(qn("w:start"), "1")

    @staticmethod
    def combine_documents_per_resolution(document_paths: List[str], output_path: str):
        """Une documentos, creando una sección numerada desde 1 para cada uno.

        El contenido siempre se inserta antes del ``sectPr`` final. Agregarlo al
        final del body produce un DOCX fuera del orden definido por WordprocessingML
        y LibreOffice puede interpretar ese XML como páginas vacías.
        """
        if not document_paths:
            raise ValueError("Se requiere al menos un documento para combinar")

        master = DocxDocument(document_paths[0])
        body = master.element.body
        for element in list(body):
            body.remove(element)

        for index, path in enumerate(document_paths):
            source = DocxDocument(path)
            copied_elements = []
            for element in source.element.body:
                if not element.tag.endswith("}sectPr"):
                    copied_element = deepcopy(element)
                    body.append(copied_element)
                    copied_elements.append(copied_element)

            section_properties = deepcopy(source.element.body.sectPr)
            DocumentGenerationService._restart_section_page_numbering(section_properties)
            if index < len(document_paths) - 1:
                section_type = section_properties.find(qn("w:type"))
                if section_type is None:
                    section_type = OxmlElement("w:type")
                    section_properties.insert(0, section_type)
                section_type.set(qn("w:val"), "nextPage")

                paragraph = next(
                    (
                        element
                        for element in reversed(copied_elements)
                        if element.tag == qn("w:p")
                    ),
                    None,
                )
                if paragraph is None:
                    paragraph = OxmlElement("w:p")
                    body.append(paragraph)
                paragraph_properties = paragraph.find(qn("w:pPr"))
                if paragraph_properties is None:
                    paragraph_properties = OxmlElement("w:pPr")
                    paragraph.insert(0, paragraph_properties)
                paragraph_properties.append(section_properties)
            else:
                body.append(section_properties)

        master.save(output_path)

    @staticmethod
    def convert_to_pdf(docx_path: str, pdf_path: str):
        """
        Convierte un archivo DOCX a PDF.
        
        Args:
            docx_path: Ruta al archivo DOCX
            pdf_path: Ruta donde guardar el archivo PDF
        """
        try:
            # Para la conversión a PDF, en un entorno real se usaría una biblioteca como python-docx2pdf
            # o se enviaría a un servicio externo. Por simplicidad, aquí solo copiamos el archivo
            # como si fuera una conversión simulada.
            
            # En un entorno real, se usaría algo como:
            # from docx2pdf import convert
            # convert(docx_path, pdf_path.replace('.pdf', '.temp.docx'))
            
            # Por ahora, simplemente creamos un archivo PDF vacío como placeholder
            # para que el sistema funcione mientras se implementa la verdadera conversión
            with open(pdf_path, 'w') as f:
                f.write(f"PDF conversion placeholder for {docx_path}")
        except Exception as e:
            raise Exception(f"Error al convertir a PDF: {str(e)}")

    @staticmethod
    def generate_batch(templates: List[Dict[str, str]], processes: List[Dict[str, Any]], output_dir: str) -> str:
        """
        Genera un lote de documentos a partir de plantillas y datos de procesos.
        
        Args:
            templates: Lista de diccionarios con información de las plantillas
            processes: Lista de diccionarios con datos de los procesos
            output_dir: Directorio donde guardar los documentos generados
            
        Returns:
            Ruta al archivo ZIP que contiene todos los documentos generados
        """
        try:
            # Crear directorio de salida si no existe
            os.makedirs(output_dir, exist_ok=True)
            
            # Crear un archivo ZIP para almacenar todos los documentos generados
            zip_filename = os.path.join(output_dir, f"batch_documents_{len(processes)}.zip")
            
            with zipfile.ZipFile(zip_filename, 'w') as zip_file:
                for i, process_data in enumerate(processes):
                    # Para cada proceso, generar un documento con cada plantilla
                    for j, template_info in enumerate(templates):
                        template_path = template_info['path']
                        
                        # Crear nombre de archivo único para el documento generado
                        output_filename = f"documento_{process_data.get('cliente_identificacion', 'desconocido')}_{i}_{j}.docx"
                        output_path = os.path.join(output_dir, output_filename)
                        
                        # Renderizar la plantilla con los datos del proceso
                        DocumentGenerationService.render_template(
                            template_path, 
                            process_data, 
                            output_path
                        )
                        
                        # Agregar el documento generado al archivo ZIP
                        zip_file.write(output_path, output_filename)
                        
                        # Opcional: eliminar el archivo temporal después de agregarlo al ZIP
                        # os.remove(output_path)
            
            return zip_filename
        except Exception as e:
            raise Exception(f"Error al generar lote de documentos: {str(e)}")

    @staticmethod
    def extract_variables_from_template(template_path: str) -> List[str]:
        """
        Extrae las variables disponibles en una plantilla DOCX.
        
        Args:
            template_path: Ruta a la plantilla DOCX
            
        Returns:
            Lista de nombres de variables encontradas en la plantilla
        """
        try:
            # Cargar la plantilla sin renderizarla
            doc = DocxTemplate(template_path)
            
            # Extraer todas las variables del documento
            # DocXTpl usa Jinja2 como motor de plantillas, así que buscamos patrones de variables
            variables = set()
            
            # Extraer variables de los párrafos
            for paragraph in doc.doc.paragraphs:
                # Buscar patrones de variables de Jinja2 {{ variable }}
                import re
                matches = re.findall(r'\{\{(\s*[a-zA-Z_][a-zA-Z0-9_]*\s*)\}\}', paragraph.text)
                for match in matches:
                    var_name = match.strip()
                    if var_name:
                        variables.add(var_name)
            
            # Extraer variables de tablas
            for table in doc.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = re.findall(r'\{\{(\s*[a-zA-Z_][a-zA-Z0-9_]*\s*)\}\}', paragraph.text)
                            for match in matches:
                                var_name = match.strip()
                                if var_name:
                                    variables.add(var_name)
            
            return list(variables)
        except Exception as e:
            raise Exception(f"Error al extraer variables de la plantilla: {str(e)}")

    @staticmethod
    def validate_template_data(template_path: str, data: Dict[str, Any]) -> Dict[str, List[str]]:
        """
        Valida que los datos proporcionados coincidan con las variables de la plantilla.
        
        Args:
            template_path: Ruta a la plantilla DOCX
            data: Datos a validar
            
        Returns:
            Diccionario con variables faltantes y variables adicionales
        """
        try:
            template_vars = set(DocumentGenerationService.extract_variables_from_template(template_path))
            data_keys = set(data.keys())
            
            missing_vars = list(template_vars - data_keys)
            extra_vars = list(data_keys - template_vars)
            
            return {
                'missing': missing_vars,
                'extra': extra_vars
            }
        except Exception as e:
            raise Exception(f"Error al validar datos de la plantilla: {str(e)}")
