from docx import Document
from docx.oxml.ns import qn

from app.services.document_generation_service import DocumentGenerationService


def _create_document(path, text):
    document = Document()
    document.add_paragraph(text)
    document.save(path)


def test_combine_documents_creates_numbered_section_per_resolution(tmp_path):
    sources = []
    for index in range(3):
        path = tmp_path / f"resolution_{index}.docx"
        _create_document(path, f"Resolución {index}")
        sources.append(str(path))

    output = tmp_path / "combined.docx"
    DocumentGenerationService.combine_documents_per_resolution(sources, str(output))

    combined = Document(output)
    assert len(combined.sections) == 3
    assert [paragraph.text for paragraph in combined.paragraphs if paragraph.text] == [
        "Resolución 0",
        "Resolución 1",
        "Resolución 2",
    ]
    for section in combined.sections:
        page_number_type = section._sectPr.find(qn("w:pgNumType"))
        assert page_number_type is not None
        assert page_number_type.get(qn("w:start")) == "1"


def test_combined_document_keeps_final_section_properties_last(tmp_path):
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    output = tmp_path / "combined.docx"
    _create_document(first, "Primera")
    _create_document(second, "Segunda")

    DocumentGenerationService.combine_documents_per_resolution(
        [str(first), str(second)], str(output)
    )

    combined = Document(output)
    body_children = list(combined.element.body)
    assert body_children[-1].tag == qn("w:sectPr")


def test_render_template_removes_empty_paragraph_before_project_signature(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "rendered.docx"
    document = Document()
    document.add_paragraph("Firma")
    document.add_paragraph("")
    document.add_paragraph("Proyecto:")
    document.add_paragraph("Responsable")
    document.save(template)

    DocumentGenerationService.render_template(str(template), {}, str(output))

    rendered = Document(output)
    project_index = next(
        index
        for index, paragraph in enumerate(rendered.paragraphs)
        if paragraph.text == "Proyecto:"
    )
    assert rendered.paragraphs[project_index - 1].text == "Firma"
    for paragraph in rendered.paragraphs[project_index:project_index + 2]:
        assert paragraph.paragraph_format.space_before.pt == 0
        assert paragraph.paragraph_format.space_after.pt == 0
        assert paragraph.paragraph_format.line_spacing == 1.0
